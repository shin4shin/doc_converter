#!/usr/bin/env python3
"""
문서 변환기 GUI — 단일 파일 버전
별도 파일 없이 이 파일 하나만으로 실행됩니다.
실행: python doc_converter_gui.py
"""

# ════════════════════════════════════════════
# 변환 엔진 (doc_converter 내장)
# ════════════════════════════════════════════
"""
문서 형식 변환기 (HWP / HWPX ↔ DOCX ↔ PDF)
지원 변환:
  - HWP  → DOCX, PDF, TXT
  - HWPX → DOCX, PDF, TXT
  - DOCX → PDF, TXT
  - PDF  → TXT, DOCX(텍스트 기반)
"""

import sys
import os
import shutil
import argparse
import subprocess
import tempfile
import zipfile
import threading
from pathlib import Path
import tkinter as tk
from tkinter import ttk, filedialog, messagebox


# ─────────────────────────────────────────────
# 유틸리티
# ─────────────────────────────────────────────

def check_dependencies():
    """필수 의존성 확인"""
    deps = {}

    # LibreOffice
    try:
        _find_libreoffice()
        deps["libreoffice"] = True
    except FileNotFoundError:
        deps["libreoffice"] = False

    # pyhwp (hwp5txt)
    deps["pyhwp"] = shutil.which("hwp5txt") is not None

    # Python 라이브러리
    try:
        import docx  # noqa
        deps["python-docx"] = True
    except ImportError:
        deps["python-docx"] = False

    try:
        import pypdf  # noqa
        deps["pypdf"] = True
    except ImportError:
        deps["pypdf"] = False

    try:
        import pdfplumber  # noqa
        deps["pdfplumber"] = True
    except ImportError:
        deps["pdfplumber"] = False

    return deps


def run_cmd(cmd, **kwargs):
    """서브프로세스 실행 후 결과 반환"""
    result = subprocess.run(cmd, capture_output=True, text=True, **kwargs)
    return result


# Windows LibreOffice 기본 설치 경로 후보
_LO_WIN_PATHS = [
    r"C:\Program Files\LibreOffice\program\soffice.exe",
    r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
    r"C:\Program Files\LibreOffice 7\program\soffice.exe",
    r"C:\Program Files\LibreOffice 24\program\soffice.exe",
]

def _find_libreoffice() -> str:
    """LibreOffice 실행 파일 경로 반환"""
    import sys
    # 환경변수/PATH 에서 먼저 찾기
    for cmd in ("libreoffice", "soffice"):
        found = shutil.which(cmd)
        if found:
            return found
    # Windows 기본 설치 경로 탐색
    if sys.platform == "win32":
        import glob
        # 버전 와일드카드로 탐색
        patterns = [
            r"C:\Program Files\LibreOffice*\program\soffice.exe",
            r"C:\Program Files (x86)\LibreOffice*\program\soffice.exe",
        ]
        for pattern in patterns:
            matches = glob.glob(pattern)
            if matches:
                return matches[0]
        for p in _LO_WIN_PATHS:
            if os.path.exists(p):
                return p
    raise FileNotFoundError(
        "LibreOffice 를 찾을 수 없습니다.\n"
        "설치 확인: https://www.libreoffice.org/download/download/\n"
        "설치 후에도 이 오류가 나오면 LibreOffice 설치 경로를 알려주세요."
    )


def lo_convert(src: Path, output_dir: Path, target_fmt: str) -> Path:
    """LibreOffice를 이용해 파일 변환"""
    lo_exe = _find_libreoffice()
    result = run_cmd(
        [lo_exe, "--headless", "--convert-to", target_fmt,
         "--outdir", str(output_dir), str(src)]
    )
    if result.returncode != 0:
        raise RuntimeError(f"LibreOffice 변환 실패:\n{result.stderr}")

    # 생성된 출력 파일 찾기
    stem = src.stem
    out_file = output_dir / f"{stem}.{target_fmt}"
    if not out_file.exists():
        matches = list(output_dir.glob(f"*.{target_fmt}"))
        if matches:
            out_file = matches[0]
        else:
            raise FileNotFoundError(f"변환된 파일을 찾을 수 없습니다: {out_file}")
    return out_file


# ─────────────────────────────────────────────
# HWP 변환
# ─────────────────────────────────────────────

def hwp_to_txt(src: Path, dst: Path):
    """HWP → TXT (pyhwp 사용)"""
    result = run_cmd(["hwp5txt", str(src)])
    if result.returncode != 0:
        raise RuntimeError(f"hwp5txt 실패:\n{result.stderr}")
    dst.write_text(result.stdout, encoding="utf-8")
    print(f"  [HWP→TXT] 텍스트 추출 완료")


def hwp_to_pdf(src: Path, dst: Path):
    """HWP → PDF (LibreOffice 사용)"""
    with tempfile.TemporaryDirectory() as tmpdir:
        out = lo_convert(src, Path(tmpdir), "pdf")
        shutil.copy2(out, dst)
    print(f"  [HWP→PDF] 변환 완료")


def hwp_to_docx(src: Path, dst: Path):
    """HWP → DOCX (LibreOffice 사용)"""
    with tempfile.TemporaryDirectory() as tmpdir:
        out = lo_convert(src, Path(tmpdir), "docx")
        shutil.copy2(out, dst)
    print(f"  [HWP→DOCX] 변환 완료")




# ─────────────────────────────────────────────
# HWPX 변환 (ZIP+XML 기반)
# ─────────────────────────────────────────────

def _hwpx_extract_text(src: Path) -> str:
    """HWPX 파일에서 텍스트 추출 (XML 직접 파싱)"""
    from lxml import etree

    NS_P = "urn:schemas-hancom-com:office:hwpx:para:1.0"
    lines = []

    try:
        with zipfile.ZipFile(str(src), "r") as z:
            names = z.namelist()
            # section 파일들 찾기 (section0.xml, section1.xml ...)
            section_files = sorted([n for n in names if "section" in n.lower() and n.endswith(".xml")])
            if not section_files:
                # content.hpf 등 다른 XML 파일에서 시도
                section_files = [n for n in names if n.endswith(".xml")]

            for sf in section_files:
                try:
                    xml_data = z.read(sf)
                    root = etree.fromstring(xml_data)
                    # hp:T 태그에서 텍스트 추출
                    for elem in root.iter(f"{{{NS_P}}}T"):
                        if elem.text:
                            lines.append(elem.text)
                    # 단락 구분
                    for elem in root.iter(f"{{{NS_P}}}P"):
                        lines.append("")
                except Exception:
                    continue
    except zipfile.BadZipFile:
        raise RuntimeError(f"올바른 HWPX 파일이 아닙니다: {src}")

    return "\n".join(lines).strip()


def hwpx_to_txt(src: Path, dst: Path):
    """HWPX → TXT (XML 직접 파싱)"""
    text = _hwpx_extract_text(src)
    if not text:
        raise RuntimeError("텍스트를 추출하지 못했습니다. 파일 구조를 확인해주세요.")
    dst.write_text(text, encoding="utf-8")
    print(f"  [HWPX→TXT] 텍스트 추출 완료")


def hwpx_to_pdf(src: Path, dst: Path):
    """HWPX → PDF (LibreOffice 시도 → XML 텍스트 폴백)"""
    with tempfile.TemporaryDirectory() as tmpdir:
        try:
            out = lo_convert(src, Path(tmpdir), "pdf")
            shutil.copy2(out, dst)
            print(f"  [HWPX→PDF] LibreOffice 변환 완료")
            return
        except Exception:
            pass

    # LibreOffice 실패 시 텍스트 → PDF (reportlab)
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
        from reportlab.lib.styles import getSampleStyleSheet
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
        import os

        text = _hwpx_extract_text(src)
        doc = SimpleDocTemplate(str(dst), pagesize=A4)
        styles = getSampleStyleSheet()

        # 한글 폰트 등록 시도
        font_paths = [
            "/usr/share/fonts/truetype/nanum/NanumGothic.ttf",
            "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
        ]
        for fp in font_paths:
            if os.path.exists(fp):
                pdfmetrics.registerFont(TTFont("MyFont", fp))
                styles["Normal"].fontName = "MyFont"
                break

        story = []
        for line in text.split("\n"):
            story.append(Paragraph(line or " ", styles["Normal"]))
            story.append(Spacer(1, 4))
        doc.build(story)
        print(f"  [HWPX→PDF] 텍스트 기반 PDF 생성 완료")
    except ImportError:
        raise RuntimeError("reportlab 을 설치하세요: pip install reportlab")


def hwpx_to_docx(src: Path, dst: Path):
    """HWPX → DOCX (LibreOffice 시도 → 텍스트 기반 폴백)"""
    with tempfile.TemporaryDirectory() as tmpdir:
        try:
            out = lo_convert(src, Path(tmpdir), "docx")
            shutil.copy2(out, dst)
            print(f"  [HWPX→DOCX] LibreOffice 변환 완료")
            return
        except Exception:
            pass

    # LibreOffice 실패 시 텍스트 → DOCX
    try:
        from docx import Document as DocxDocument
        text = _hwpx_extract_text(src)
        doc = DocxDocument()
        for line in text.split("\n"):
            doc.add_paragraph(line)
        doc.save(str(dst))
        print(f"  [HWPX→DOCX] 텍스트 기반 DOCX 생성 완료")
    except ImportError:
        raise RuntimeError("python-docx 를 설치하세요: pip install python-docx")

# ─────────────────────────────────────────────
# DOCX 변환
# ─────────────────────────────────────────────

def docx_to_pdf(src: Path, dst: Path):
    """DOCX → PDF (LibreOffice 사용)"""
    with tempfile.TemporaryDirectory() as tmpdir:
        out = lo_convert(src, Path(tmpdir), "pdf")
        shutil.copy2(out, dst)
    print(f"  [DOCX→PDF] 변환 완료")


def docx_to_txt(src: Path, dst: Path):
    """DOCX → TXT (python-docx 사용)"""
    try:
        from docx import Document
    except ImportError:
        raise RuntimeError("python-docx 가 설치되지 않았습니다: pip install python-docx")

    doc = Document(str(src))
    lines = []
    for para in doc.paragraphs:
        lines.append(para.text)
    # 표 안의 텍스트도 추출
    for table in doc.tables:
        for row in table.rows:
            row_texts = [cell.text.strip() for cell in row.cells]
            lines.append("\t".join(row_texts))

    dst.write_text("\n".join(lines), encoding="utf-8")
    print(f"  [DOCX→TXT] 텍스트 추출 완료")


def _escape_xml(text: str) -> str:
    """XML 특수문자 이스케이프"""
    return (text
        .replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
        .replace('"', "&quot;"))


def docx_to_hwpx_file(src: Path, dst: Path):
    """DOCX → HWPX (ZIP+XML 직접 생성, 텍스트·표 보존)"""
    try:
        from docx import Document
    except ImportError:
        raise RuntimeError("python-docx 를 설치하세요: pip install python-docx")

    doc = Document(str(src))

    NS_P = "urn:schemas-hancom-com:office:hwpx:para:1.0"
    NS_H = "urn:schemas-hancom-com:office:hwpx:1.0"

    paras = []

    # 단락 추출 (스타일 정보 포함)
    for para in doc.paragraphs:
        text = _escape_xml(para.text)
        style = para.style.name if para.style else ""
        if "Heading" in style:
            paras.append(f'  <hp:P><hp:Run><hp:T><![CDATA[]]></hp:T></hp:Run></hp:P>')
            paras.append(f'  <hp:P><hp:Run><hp:T>{text}</hp:T></hp:Run></hp:P>')
        else:
            paras.append(f'  <hp:P><hp:Run><hp:T>{text}</hp:T></hp:Run></hp:P>')

    # 표 추출
    for table in doc.tables:
        paras.append('  <hp:P><hp:Run><hp:T></hp:T></hp:Run></hp:P>')
        for row in table.rows:
            row_text = " | ".join(_escape_xml(cell.text.strip()) for cell in row.cells)
            paras.append(f'  <hp:P><hp:Run><hp:T>{row_text}</hp:T></hp:Run></hp:P>')
        paras.append('  <hp:P><hp:Run><hp:T></hp:T></hp:Run></hp:P>')

    section_xml = (
        '<?xml version="1.0" encoding="UTF-8"?>\n'
        f'<hp:Section xmlns:hp="{NS_P}">\n'
        + "\n".join(paras)
        + "\n</hp:Section>"
    ).encode("utf-8")

    content_hpf = (
        '<?xml version="1.0" encoding="UTF-8"?>\n'
        f'<hh:HWPXDocument xmlns:hh="{NS_H}" version="1.0.0.0">\n'
        '  <hh:Body>\n'
        '    <hh:SectionList>\n'
        '      <hh:SectionRef href="Contents/section0.xml"/>\n'
        '    </hh:SectionList>\n'
        '  </hh:Body>\n'
        '</hh:HWPXDocument>'
    ).encode("utf-8")

    content_types = (
        '<?xml version="1.0" encoding="UTF-8"?>\n'
        '<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">'
        '<Default Extension="xml" ContentType="application/xml"/>'
        '<Override PartName="/Contents/content.hpf"'
        ' ContentType="application/vnd.hancom.hwpx"/>'
        '</Types>'
    ).encode("utf-8")

    with zipfile.ZipFile(str(dst), "w", zipfile.ZIP_DEFLATED) as z:
        z.writestr("[Content_Types].xml", content_types)
        z.writestr("Contents/content.hpf", content_hpf)
        z.writestr("Contents/section0.xml", section_xml)

    print(f"  [DOCX→HWPX] 변환 완료 (단락 {len([p for p in doc.paragraphs if p.text])}개)")


def docx_to_hwp(src: Path, dst: Path):
    """DOCX → HWPX 로 저장 (HWP 바이너리 직접 생성 불가, HWPX로 대체)"""
    # 출력 확장자를 .hwpx 로 변경
    hwpx_dst = dst.with_suffix(".hwpx")
    docx_to_hwpx_file(src, hwpx_dst)
    # 요청이 .hwp 였으면 안내
    if dst.suffix.lower() == ".hwp":
        print(f"  ℹ️  HWP 바이너리 포맷은 직접 생성이 불가하여 HWPX 로 저장했습니다: {hwpx_dst.name}")
        # dst 도 hwpx 로 복사
        shutil.copy2(hwpx_dst, hwpx_dst)  # 이미 저장됨
    print(f"  ✅ 저장 완료: {hwpx_dst}")


# ─────────────────────────────────────────────
# PDF 변환
# ─────────────────────────────────────────────

def pdf_to_txt(src: Path, dst: Path):
    """PDF → TXT (pdfplumber 우선, pypdf 폴백)"""
    try:
        import pdfplumber
        text_parts = []
        with pdfplumber.open(str(src)) as pdf:
            for i, page in enumerate(pdf.pages, 1):
                t = page.extract_text()
                if t:
                    text_parts.append(f"=== 페이지 {i} ===\n{t}")
        dst.write_text("\n\n".join(text_parts), encoding="utf-8")
        print(f"  [PDF→TXT] {len(text_parts)}페이지 추출 완료")
        return
    except ImportError:
        pass

    try:
        from pypdf import PdfReader
        reader = PdfReader(str(src))
        text_parts = []
        for i, page in enumerate(reader.pages, 1):
            t = page.extract_text()
            if t:
                text_parts.append(f"=== 페이지 {i} ===\n{t}")
        dst.write_text("\n\n".join(text_parts), encoding="utf-8")
        print(f"  [PDF→TXT] {len(text_parts)}페이지 추출 완료")
    except ImportError:
        raise RuntimeError("pypdf 또는 pdfplumber 를 설치하세요: pip install pypdf pdfplumber")


def pdf_to_docx(src: Path, dst: Path):
    """PDF → DOCX (텍스트 기반, LibreOffice 사용)"""
    with tempfile.TemporaryDirectory() as tmpdir:
        try:
            out = lo_convert(src, Path(tmpdir), "docx")
            shutil.copy2(out, dst)
            print(f"  [PDF→DOCX] 변환 완료")
        except Exception as e:
            raise RuntimeError(
                f"PDF→DOCX 변환 실패: {e}\n"
                "스캔 PDF의 경우 OCR 소프트웨어가 필요합니다."
            )


def pdf_to_hwp(src: Path, dst: Path):
    """PDF → HWPX (PDF → DOCX → HWPX 경유)"""
    with tempfile.TemporaryDirectory() as tmpdir:
        docx_tmp = Path(tmpdir) / (src.stem + ".docx")
        pdf_to_docx(src, docx_tmp)
        hwpx_dst = dst.with_suffix(".hwpx")
        docx_to_hwpx_file(docx_tmp, hwpx_dst)
        print(f"  ℹ️  HWP 대신 HWPX 로 저장됩니다: {hwpx_dst.name}")


def pdf_to_hwpx(src: Path, dst: Path):
    """PDF → HWPX (PDF → DOCX → HWPX 경유)"""
    with tempfile.TemporaryDirectory() as tmpdir:
        docx_tmp = Path(tmpdir) / (src.stem + ".docx")
        pdf_to_docx(src, docx_tmp)
        docx_to_hwpx_file(docx_tmp, dst)


# ─────────────────────────────────────────────
# 라우터
# ─────────────────────────────────────────────

CONVERTERS = {
    ("hwp",  "pdf"):  hwp_to_pdf,
    ("hwp",  "docx"): hwp_to_docx,
    ("hwp",  "txt"):  hwp_to_txt,
    ("hwpx", "pdf"):  hwpx_to_pdf,
    ("hwpx", "docx"): hwpx_to_docx,
    ("hwpx", "txt"):  hwpx_to_txt,
    ("docx", "pdf"):  docx_to_pdf,
    ("docx", "txt"):  docx_to_txt,
    ("docx", "hwp"):  docx_to_hwp,
    ("docx", "hwpx"): docx_to_hwpx_file,
    ("pdf",  "txt"):  pdf_to_txt,
    ("pdf",  "docx"): pdf_to_docx,
    ("pdf",  "hwp"):  pdf_to_hwp,
    ("pdf",  "hwpx"): pdf_to_hwpx,
}


def convert(src_path: str, target_fmt: str, output_path: str = None):
    src = Path(src_path).resolve()
    if not src.exists():
        print(f"오류: 파일을 찾을 수 없습니다 → {src}")
        sys.exit(0)

    src_fmt = src.suffix.lstrip(".").lower()
    tgt_fmt = target_fmt.lower().lstrip(".")

    key = (src_fmt, tgt_fmt)
    if key not in CONVERTERS:
        supported = "\n  ".join(f"{a} → {b}" for a, b in sorted(CONVERTERS))
        print(f"오류: '{src_fmt}' → '{tgt_fmt}' 변환은 지원되지 않습니다.")
        print(f"지원 변환 목록:\n  {supported}")
        sys.exit(0)

    if output_path:
        dst = Path(output_path).resolve()
    else:
        dst = src.with_suffix(f".{tgt_fmt}")

    dst.parent.mkdir(parents=True, exist_ok=True)

    print(f"변환 시작: {src.name}  →  {dst.name}")
    try:
        CONVERTERS[key](src, dst)
        print(f"✅ 완료: {dst}")
    except Exception as e:
        print(f"❌ 변환 실패: {e}")
        sys.exit(0)


# ─────────────────────────────────────────────
# CLI
# ─────────────────────────────────────────────

def main():
    parser = argparse.ArgumentParser(
        description="문서 형식 변환기 (HWP / DOCX / PDF)",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
사용 예시:
  python doc_converter.py report.hwp pdf
  python doc_converter.py report.hwp docx
  python doc_converter.py document.docx pdf
  python doc_converter.py document.docx txt
  python doc_converter.py file.pdf docx
  python doc_converter.py file.pdf txt --output result.txt

지원 변환:
  HWP  → PDF, DOCX, TXT
  HWPX → PDF, DOCX, TXT
  DOCX → PDF, TXT
  PDF  → TXT, DOCX
        """,
    )
    parser.add_argument("input",  nargs="?", help="변환할 원본 파일 경로")
    parser.add_argument("format", nargs="?", help="변환 대상 형식 (예: pdf, docx, txt, hwp)")
    parser.add_argument("-o", "--output", help="출력 파일 경로 (기본: 원본과 같은 폴더)", default=None)
    parser.add_argument("--check", action="store_true", help="의존성 확인만 실행")


    args = parser.parse_args()

    if args.check:
        print("=== 의존성 확인 ===")
        deps = check_dependencies()
        for name, ok in deps.items():
            status = "✅" if ok else "❌"
            print(f"  {status}  {name}")
        print()
        missing = [n for n, ok in deps.items() if not ok]
        if missing:
            print("미설치 항목:")
            install_hints = {
                "libreoffice": "sudo apt install libreoffice",
                "pyhwp":       "pip install pyhwp",
                "python-docx": "pip install python-docx",
                "pypdf":       "pip install pypdf",
                "pdfplumber":  "pip install pdfplumber",
            }
            for m in missing:
                print(f"  {m}: {install_hints.get(m, '설치 필요')}")
        else:
            print("모든 의존성이 설치되어 있습니다.")
        return

    if not args.input or not args.format:
        parser.print_help()
        print()
        print("사용법: python doc_converter.py <파일경로> <변환형식>")
        print("예시:   python doc_converter.py 문서.hwp pdf")
        sys.exit(0)

    convert(args.input, args.format, args.output)



# ════════════════════════════════════════════
# GUI
# ════════════════════════════════════════════
# ── 지원 변환 테이블 ──────────────────────────────
SUPPORT = {
    "hwp":  ["pdf", "docx", "txt"],
    "hwpx": ["pdf", "docx", "txt"],
    "docx": ["pdf", "txt", "hwp", "hwpx"],
    "pdf":  ["txt", "docx", "hwp", "hwpx"],
}

EXT_LABEL = {
    "hwp":  "HWP  (한글 97~2018)",
    "hwpx": "HWPX (한글 2019+)",
    "docx": "DOCX (Word)",
    "pdf":  "PDF",
    "txt":  "TXT  (텍스트)",
}

# ── 색상 팔레트 ───────────────────────────────────
BG       = "#1e1e2e"
PANEL    = "#2a2a3e"
ACCENT   = "#7c6af7"
ACCENT2  = "#56cfb2"
TEXT     = "#e0e0f0"
SUBTEXT  = "#888aaa"
SUCCESS  = "#56cfb2"
ERROR    = "#f76a6a"
BORDER   = "#3a3a55"
BTN_FG   = "#ffffff"
ENTRY_BG = "#252538"


class App(tk.Tk):
    def __init__(self):
        super().__init__()
        self.title("문서 변환기 v1.1.0")
        self.resizable(False, False)
        self.configure(bg=BG)

        # 상태
        self.src_path   = tk.StringVar()
        self.src_ext    = tk.StringVar()   # 감지된 입력 확장자
        self.tgt_fmt    = tk.StringVar()
        self.out_path   = tk.StringVar()
        self.status_msg = tk.StringVar(value="파일을 선택해주세요.")
        self.converting = False

        self._build_ui()
        self._center()

    # ── UI 구성 ──────────────────────────────────

    def _build_ui(self):
        px = dict(padx=24)

        # ── 타이틀 바 ──
        title_bar = tk.Frame(self, bg=ACCENT, height=4)
        title_bar.pack(fill="x")

        header = tk.Frame(self, bg=BG)
        header.pack(fill="x", padx=24, pady=(20, 4))

        tk.Label(header, text="문서 변환기", font=("Malgun Gothic", 20, "bold"),
                 bg=BG, fg=TEXT).pack(side="left")
        tk.Label(header, text="HWP · HWPX · DOCX · PDF · TXT",
                 font=("Malgun Gothic", 9), bg=BG, fg=SUBTEXT).pack(side="left", padx=(12, 0), pady=(6, 0))

        self._divider()

        # ── 섹션 1: 입력 파일 ──
        self._section_label("① 변환할 파일 선택")
        row1 = tk.Frame(self, bg=BG)
        row1.pack(fill="x", **px, pady=(0, 12))

        self.entry_src = self._entry(row1, self.src_path, placeholder="파일을 선택하거나 경로를 직접 입력하세요")
        self.entry_src.pack(side="left", fill="x", expand=True)
        self.entry_src.bind("<FocusOut>", lambda e: self._on_src_changed())
        self.entry_src.bind("<Return>",   lambda e: self._on_src_changed())

        self._btn(row1, "열기", self._browse_src, accent=True).pack(side="left", padx=(8, 0))

        # 감지된 형식 뱃지
        self.badge_frame = tk.Frame(self, bg=BG)
        self.badge_frame.pack(fill="x", padx=24, pady=(0, 8))
        self.badge_lbl = tk.Label(self.badge_frame, text="", font=("Malgun Gothic", 9),
                                  bg=BG, fg=ACCENT2)
        self.badge_lbl.pack(side="left")

        self._divider()

        # ── 섹션 2: 변환 형식 ──
        self._section_label("② 변환할 형식 선택")
        self.fmt_frame = tk.Frame(self, bg=BG)
        self.fmt_frame.pack(fill="x", padx=24, pady=(0, 16))
        self.fmt_btns = {}   # fmt -> Radiobutton

        # 처음엔 전부 비활성
        self._rebuild_fmt_buttons(available=[])

        self._divider()

        # ── 섹션 3: 출력 경로 ──
        self._section_label("③ 저장 경로 (선택)")
        row3 = tk.Frame(self, bg=BG)
        row3.pack(fill="x", **px, pady=(0, 4))

        self.entry_out = self._entry(row3, self.out_path,
                                     placeholder="비워두면 원본 파일과 같은 폴더에 저장됩니다")
        self.entry_out.pack(side="left", fill="x", expand=True)
        self._btn(row3, "저장 위치", self._browse_dst).pack(side="left", padx=(8, 0))

        self._divider()

        # ── 변환 버튼 ──
        btn_row = tk.Frame(self, bg=BG)
        btn_row.pack(fill="x", padx=24, pady=(4, 0))
        self.convert_btn = self._btn(btn_row, "변환 시작  ▶", self._start_convert,
                                     accent=True, big=True)
        self.convert_btn.pack(fill="x")

        # ── 진행 바 ──
        self.progress = ttk.Progressbar(self, mode="indeterminate", length=400)
        style = ttk.Style(self)
        style.theme_use("clam")
        style.configure("TProgressbar", troughcolor=PANEL, background=ACCENT,
                        bordercolor=BG, lightcolor=ACCENT, darkcolor=ACCENT)
        self.progress.pack(fill="x", padx=24, pady=(10, 0))

        # ── 상태 메시지 ──
        status_row = tk.Frame(self, bg=BG)
        status_row.pack(fill="x", padx=24, pady=(8, 20))
        self.status_icon = tk.Label(status_row, text="●", font=("Malgun Gothic", 10),
                                    bg=BG, fg=SUBTEXT)
        self.status_icon.pack(side="left")
        tk.Label(status_row, textvariable=self.status_msg,
                 font=("Malgun Gothic", 10), bg=BG, fg=SUBTEXT,
                 wraplength=440, justify="left").pack(side="left", padx=(6, 0))

    # ── 헬퍼: 위젯 생성 ──────────────────────────

    def _divider(self):
        tk.Frame(self, bg=BORDER, height=1).pack(fill="x", padx=16, pady=8)

    def _section_label(self, text):
        tk.Label(self, text=text, font=("Malgun Gothic", 11, "bold"),
                 bg=BG, fg=TEXT).pack(anchor="w", padx=24, pady=(4, 6))

    def _entry(self, parent, var, placeholder=""):
        e = tk.Entry(parent, textvariable=var,
                     font=("Malgun Gothic", 10),
                     bg=ENTRY_BG, fg=TEXT, insertbackground=TEXT,
                     relief="flat", bd=0, highlightthickness=1,
                     highlightbackground=BORDER, highlightcolor=ACCENT)
        e.configure(width=44)
        # placeholder 구현
        if placeholder and not var.get():
            e.insert(0, placeholder)
            e.configure(fg=SUBTEXT)
            def on_focus_in(ev, _e=e, _v=var, _p=placeholder):
                if _e.get() == _p:
                    _e.delete(0, "end")
                    _e.configure(fg=TEXT)
            def on_focus_out(ev, _e=e, _v=var, _p=placeholder):
                if not _e.get():
                    _e.insert(0, _p)
                    _e.configure(fg=SUBTEXT)
            e.bind("<FocusIn>",  on_focus_in)
            e.bind("<FocusOut>", on_focus_out)
        return e

    def _btn(self, parent, text, cmd, accent=False, big=False):
        bg = ACCENT if accent else PANEL
        pad_y = 10 if big else 6
        fsize = 11 if big else 10
        b = tk.Button(parent, text=text, command=cmd,
                      font=("Malgun Gothic", fsize, "bold" if big else "normal"),
                      bg=bg, fg=BTN_FG, activebackground=ACCENT2,
                      activeforeground=BTN_FG, relief="flat", bd=0,
                      padx=14, pady=pad_y, cursor="hand2")
        b.bind("<Enter>", lambda e, w=b, _bg=bg: w.configure(bg=ACCENT2))
        b.bind("<Leave>", lambda e, w=b, _bg=bg: w.configure(bg=_bg))
        return b

    def _rebuild_fmt_buttons(self, available: list):
        for w in self.fmt_frame.winfo_children():
            w.destroy()
        self.fmt_btns.clear()

        all_fmts = ["pdf", "docx", "txt", "hwp", "hwpx"]
        for fmt in all_fmts:
            active = fmt in available
            bg_col  = PANEL  if active else BG
            fg_col  = TEXT   if active else BORDER
            bdr_col = ACCENT if active else BORDER

            btn = tk.Radiobutton(
                self.fmt_frame, text=EXT_LABEL.get(fmt, fmt).split("(")[0].strip(),
                variable=self.tgt_fmt, value=fmt,
                font=("Malgun Gothic", 10),
                bg=bg_col, fg=fg_col, selectcolor=ACCENT,
                activebackground=bg_col, activeforeground=TEXT,
                relief="flat", bd=0,
                indicatoron=False,
                padx=14, pady=8,
                highlightthickness=1,
                highlightbackground=bdr_col,
                state="normal" if active else "disabled",
                cursor="hand2" if active else "arrow",
            )
            btn.pack(side="left", padx=(0, 8))
            self.fmt_btns[fmt] = btn

        if available and self.tgt_fmt.get() not in available:
            self.tgt_fmt.set(available[0])

    # ── 이벤트 핸들러 ─────────────────────────────

    def _browse_src(self):
        filetypes = [
            ("지원 문서", "*.hwp *.hwpx *.docx *.pdf"),
            ("HWP 파일", "*.hwp"),
            ("HWPX 파일", "*.hwpx"),
            ("Word 문서", "*.docx"),
            ("PDF 파일", "*.pdf"),
            ("모든 파일", "*.*"),
        ]
        path = filedialog.askopenfilename(filetypes=filetypes)
        if path:
            self.src_path.set(path)
            self.entry_src.configure(fg=TEXT)
            self._on_src_changed()

    def _browse_dst(self):
        fmt = self.tgt_fmt.get() or "pdf"
        path = filedialog.asksaveasfilename(
            defaultextension=f".{fmt}",
            filetypes=[(fmt.upper(), f"*.{fmt}"), ("모든 파일", "*.*")]
        )
        if path:
            self.out_path.set(path)
            self.entry_out.configure(fg=TEXT)

    def _on_src_changed(self):
        raw = self.src_path.get().strip()
        # placeholder 텍스트 무시
        if not raw or "파일을 선택" in raw:
            self.badge_lbl.configure(text="")
            self._rebuild_fmt_buttons(available=[])
            return

        ext = Path(raw).suffix.lstrip(".").lower()
        if ext in SUPPORT:
            self.src_ext.set(ext)
            label = EXT_LABEL.get(ext, ext)
            self.badge_lbl.configure(text=f"감지된 형식: {label}", fg=ACCENT2)
            self._rebuild_fmt_buttons(available=SUPPORT[ext])
            self._set_status("형식을 선택하고 변환을 시작하세요.", color=SUBTEXT)
        else:
            self.src_ext.set("")
            self.badge_lbl.configure(text=f"지원하지 않는 형식: .{ext}", fg=ERROR)
            self._rebuild_fmt_buttons(available=[])

    def _set_status(self, msg, color=SUBTEXT):
        self.status_msg.set(msg)
        icon_color = SUCCESS if color == SUCCESS else (ERROR if color == ERROR else SUBTEXT)
        self.status_icon.configure(fg=icon_color)

    def _start_convert(self):
        if self.converting:
            return

        src = self.src_path.get().strip()
        fmt = self.tgt_fmt.get().strip()
        out = self.out_path.get().strip()

        # 유효성 검사
        if not src or "파일을 선택" in src:
            messagebox.showwarning("알림", "변환할 파일을 선택해주세요.")
            return
        if not Path(src).exists():
            messagebox.showerror("오류", f"파일을 찾을 수 없습니다:\n{src}")
            return
        if not fmt:
            messagebox.showwarning("알림", "변환할 형식을 선택해주세요.")
            return

        # placeholder 처리
        if "저장 위치" in out or "비워두면" in out:
            out = ""

        self.converting = True
        self.convert_btn.configure(state="disabled", text="변환 중…")
        self.progress.start(12)
        self._set_status("변환 중입니다…", color=SUBTEXT)

        thread = threading.Thread(target=self._run_convert,
                                  args=(src, fmt, out or None), daemon=True)
        thread.start()

    def _run_convert(self, src, fmt, out):
        try:
            convert(src, fmt, out)
            dst = out if out else str(Path(src).with_suffix(f".{fmt}"))
            self.after(0, self._on_success, dst)
        except Exception as e:
            self.after(0, self._on_error, str(e))

    def _on_success(self, dst):
        self.converting = False
        self.progress.stop()
        self.convert_btn.configure(state="normal", text="변환 시작  ▶")
        self._set_status(f"✅ 완료: {dst}", color=SUCCESS)
        self.status_icon.configure(fg=SUCCESS)

        result = messagebox.askyesno("완료", f"변환이 완료됐습니다!\n\n{dst}\n\n저장 폴더를 여시겠습니까?")
        if result:
            folder = str(Path(dst).parent)
            if sys.platform == "win32":
                os.startfile(folder)
            elif sys.platform == "darwin":
                subprocess.run(["open", folder])
            else:
                subprocess.run(["xdg-open", folder])

    def _on_error(self, msg):
        self.converting = False
        self.progress.stop()
        self.convert_btn.configure(state="normal", text="변환 시작  ▶")
        self._set_status(f"❌ 실패: {msg}", color=ERROR)
        messagebox.showerror("변환 실패", msg)

    # ── 창 가운데 배치 ────────────────────────────

    def _center(self):
        self.update_idletasks()
        w, h = self.winfo_width(), self.winfo_height()
        sw = self.winfo_screenwidth()
        sh = self.winfo_screenheight()
        self.geometry(f"+{(sw-w)//2}+{(sh-h)//2}")


if __name__ == "__main__":
    app = App()
    app.mainloop()